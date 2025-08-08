"""
AITM Report Formatting and Export Service

This module provides comprehensive report formatting capabilities including PDF, HTML, DOCX, 
and other output formats. It handles template management, chart generation, and professional 
document layout.
"""

import asyncio
import logging
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, BinaryIO
import json
import base64
from io import BytesIO

# PDF generation
try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    logging.warning("WeasyPrint not available. PDF generation will be disabled.")

# Chart generation  
try:
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches
    import seaborn as sns
    import pandas as pd
    import numpy as np
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False
    logging.warning("Matplotlib not available. Chart generation will be limited.")

# Document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX generation will be disabled.")

from app.agents.report_generator import ReportContent, ReportFormat, ReportType
from app.core.config import get_settings

settings = get_settings()
logger = logging.getLogger(__name__)


class ChartGenerator:
    """Generate charts and visualizations for reports"""
    
    def __init__(self):
        self.chart_style = {
            'figure_size': (10, 6),
            'dpi': 300,
            'color_palette': ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', 
                            '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf'],
            'background_color': 'white',
            'grid_alpha': 0.3
        }
    
    async def generate_chart(self, chart_data: Dict[str, Any]) -> str:
        """Generate a chart and return base64 encoded image"""
        if not MATPLOTLIB_AVAILABLE:
            return self._create_placeholder_chart()
        
        chart_type = chart_data.get("type", "bar")
        title = chart_data.get("title", "Chart")
        data = chart_data.get("data", [])
        
        plt.style.use('seaborn-v0_8')
        fig, ax = plt.subplots(figsize=self.chart_style['figure_size'], 
                              dpi=self.chart_style['dpi'])
        
        try:
            if chart_type == "pie":
                await self._generate_pie_chart(ax, data, title)
            elif chart_type == "bar":
                await self._generate_bar_chart(ax, data, title)
            elif chart_type == "line":
                await self._generate_line_chart(ax, data, title)
            elif chart_type == "heatmap":
                await self._generate_heatmap_chart(ax, data, title)
            elif chart_type == "radar":
                await self._generate_radar_chart(fig, ax, data, title)
            elif chart_type == "donut":
                await self._generate_donut_chart(ax, data, title)
            else:
                await self._generate_bar_chart(ax, data, title)
            
            # Save to base64
            buffer = BytesIO()
            plt.tight_layout()
            plt.savefig(buffer, format='png', bbox_inches='tight', 
                       facecolor=self.chart_style['background_color'])
            buffer.seek(0)
            
            # Convert to base64
            chart_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
            plt.close(fig)
            
            return f"data:image/png;base64,{chart_base64}"
            
        except Exception as e:
            logger.error(f"Error generating chart: {e}")
            plt.close(fig)
            return self._create_placeholder_chart()
    
    async def _generate_pie_chart(self, ax, data: List[Dict], title: str):
        """Generate pie chart"""
        labels = [item.get("label", "Unknown") for item in data]
        values = [item.get("value", 0) for item in data]
        colors = self.chart_style['color_palette'][:len(data)]
        
        wedges, texts, autotexts = ax.pie(values, labels=labels, autopct='%1.1f%%', 
                                         colors=colors, startangle=90)
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        
        # Beautify text
        for text in texts:
            text.set_fontsize(10)
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
    
    async def _generate_bar_chart(self, ax, data: List[Dict], title: str):
        """Generate bar chart"""
        labels = [item.get("label", "Unknown") for item in data]
        values = [item.get("value", 0) for item in data]
        colors = self.chart_style['color_palette'][:len(data)]
        
        bars = ax.bar(labels, values, color=colors, alpha=0.8)
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        ax.set_xlabel('Categories', fontsize=12)
        ax.set_ylabel('Count', fontsize=12)
        ax.grid(True, alpha=self.chart_style['grid_alpha'])
        
        # Add value labels on bars
        for bar in bars:
            height = bar.get_height()
            ax.annotate(f'{int(height)}',
                       xy=(bar.get_x() + bar.get_width() / 2, height),
                       xytext=(0, 3),  # 3 points vertical offset
                       textcoords="offset points",
                       ha='center', va='bottom', fontweight='bold')
        
        # Rotate labels if needed
        if len(labels) > 5:
            plt.setp(ax.get_xticklabels(), rotation=45, ha='right')
    
    async def _generate_line_chart(self, ax, data: List[Dict], title: str):
        """Generate line chart"""
        x_values = [item.get("date", i) for i, item in enumerate(data)]
        y_values = [item.get("risks", 0) for item in data]
        
        ax.plot(x_values, y_values, marker='o', linewidth=2, 
               markersize=6, color=self.chart_style['color_palette'][0])
        ax.fill_between(x_values, y_values, alpha=0.3, 
                       color=self.chart_style['color_palette'][0])
        
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        ax.set_xlabel('Date', fontsize=12)
        ax.set_ylabel('Risk Count', fontsize=12)
        ax.grid(True, alpha=self.chart_style['grid_alpha'])
        
        # Format x-axis if dates
        if len(x_values) > 5:
            ax.tick_params(axis='x', rotation=45)
    
    async def _generate_heatmap_chart(self, ax, data: List[Dict], title: str):
        """Generate heatmap chart"""
        # Create matrix from data
        tactics = [item.get("tactic", "Unknown") for item in data]
        counts = [item.get("count", 0) for item in data]
        
        # Create a simple heatmap representation
        max_count = max(counts) if counts else 1
        normalized_counts = [c / max_count for c in counts]
        
        # Create matrix-like visualization
        matrix = np.array(normalized_counts).reshape(-1, 1)
        
        im = ax.imshow(matrix, cmap='Reds', aspect='auto')
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        ax.set_xticks([])
        ax.set_yticks(range(len(tactics)))
        ax.set_yticklabels(tactics)
        
        # Add colorbar
        plt.colorbar(im, ax=ax, label='Relative Frequency')
    
    async def _generate_radar_chart(self, fig, ax, data: List[Dict], title: str):
        """Generate radar chart"""
        frameworks = [item.get("framework", "Unknown") for item in data]
        scores = [item.get("score", 0) for item in data]
        
        # Number of variables
        N = len(frameworks)
        
        # Compute angle for each axis
        angles = [n / float(N) * 2 * np.pi for n in range(N)]
        angles += angles[:1]  # Complete the circle
        
        # Close the plot
        scores += scores[:1]
        
        # Plot
        ax = fig.add_subplot(111, projection='polar')
        ax.plot(angles, scores, 'o-', linewidth=2, 
               color=self.chart_style['color_palette'][0])
        ax.fill(angles, scores, alpha=0.25, 
               color=self.chart_style['color_palette'][0])
        
        # Add labels
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(frameworks)
        ax.set_ylim(0, 100)
        ax.set_title(title, fontsize=14, fontweight='bold', pad=30)
        ax.grid(True)
    
    async def _generate_donut_chart(self, ax, data: List[Dict], title: str):
        """Generate donut chart"""
        labels = [item.get("label", "Unknown") for item in data]
        values = [item.get("value", 0) for item in data]
        colors = data[0].get("colors", self.chart_style['color_palette'][:len(data)])
        
        # Create pie chart with hole in center
        wedges, texts, autotexts = ax.pie(values, labels=labels, autopct='%1.1f%%',
                                         colors=colors, startangle=90, 
                                         wedgeprops=dict(width=0.5))
        ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
        
        # Draw center circle to create donut effect
        centre_circle = plt.Circle((0, 0), 0.70, fc='white')
        fig = plt.gcf()
        fig.gca().add_artist(centre_circle)
    
    def _create_placeholder_chart(self) -> str:
        """Create placeholder chart when matplotlib is not available"""
        # Return a simple base64 encoded 1x1 pixel image
        pixel_data = b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\tpHYs\x00\x00\x0b\x13\x00\x00\x0b\x13\x01\x00\x9a\x9c\x18\x00\x00\x00\x0bIDATx\x9cc\xf8\x00\x00\x00\x00\x00\x01\x00\x00\x00\x00\x37\x00\x00\x00\x00IEND\xaeB`\x82'
        return f"data:image/png;base64,{base64.b64encode(pixel_data).decode('utf-8')}"


class HTMLTemplateEngine:
    """HTML template engine for report generation"""
    
    def __init__(self):
        self.base_template = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{title}}</title>
    <style>
        {{styles}}
    </style>
</head>
<body>
    <div class="report-container">
        {{content}}
    </div>
</body>
</html>
        """.strip()
        
        self.default_styles = """
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            margin: 0;
            padding: 20px;
            background-color: #f9f9f9;
        }
        .report-container {
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            overflow: hidden;
        }
        .header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 30px;
            text-align: center;
        }
        .header h1 {
            margin: 0;
            font-size: 2.5em;
            font-weight: 300;
        }
        .header .subtitle {
            margin-top: 10px;
            opacity: 0.9;
            font-size: 1.1em;
        }
        .executive-summary {
            background: #f8f9ff;
            padding: 30px;
            border-left: 4px solid #667eea;
            margin: 20px;
            border-radius: 4px;
        }
        .executive-summary h2 {
            color: #667eea;
            margin-top: 0;
        }
        .content-section {
            padding: 20px 30px;
            margin: 20px;
            border-radius: 4px;
            background: white;
            border: 1px solid #e1e5e9;
        }
        .section-title {
            color: #333;
            border-bottom: 2px solid #667eea;
            padding-bottom: 10px;
            margin-bottom: 20px;
        }
        .metrics-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin: 20px 0;
        }
        .metric-card {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 8px;
            text-align: center;
        }
        .metric-value {
            font-size: 2.5em;
            font-weight: bold;
            margin-bottom: 5px;
        }
        .metric-label {
            opacity: 0.9;
            font-size: 0.9em;
            text-transform: uppercase;
            letter-spacing: 1px;
        }
        .chart-container {
            text-align: center;
            margin: 30px 0;
            padding: 20px;
            background: #fafafa;
            border-radius: 4px;
        }
        .chart-container img {
            max-width: 100%;
            height: auto;
            border-radius: 4px;
        }
        .recommendations {
            background: #fff8e7;
            border-left: 4px solid #f59e0b;
            padding: 20px;
            margin: 20px 0;
            border-radius: 4px;
        }
        .recommendations h3 {
            color: #f59e0b;
            margin-top: 0;
        }
        .recommendation-item {
            margin: 15px 0;
            padding: 15px;
            background: white;
            border-radius: 4px;
            border: 1px solid #fbbf24;
        }
        .priority-high { border-left: 4px solid #dc2626; }
        .priority-medium { border-left: 4px solid #f59e0b; }
        .priority-low { border-left: 4px solid #10b981; }
        
        .finding-item {
            margin: 15px 0;
            padding: 15px;
            background: white;
            border-radius: 4px;
            border: 1px solid #d1d5db;
        }
        .impact-high { border-left: 4px solid #dc2626; }
        .impact-medium { border-left: 4px solid #f59e0b; }
        .impact-low { border-left: 4px solid #10b981; }
        
        .footer {
            background: #1f2937;
            color: white;
            padding: 30px;
            text-align: center;
            font-size: 0.9em;
        }
        .footer .timestamp {
            opacity: 0.7;
        }
        
        @media print {
            body { background: white; padding: 0; }
            .report-container { box-shadow: none; }
        }
        """
    
    async def render_report(self, content: ReportContent) -> str:
        """Render report content to HTML"""
        html_content = await self._build_html_content(content)
        
        return self.base_template.replace("{{title}}", content.title) \
                                .replace("{{styles}}", self.default_styles) \
                                .replace("{{content}}", html_content)
    
    async def _build_html_content(self, content: ReportContent) -> str:
        """Build the main HTML content"""
        html_parts = []
        
        # Header
        html_parts.append(f"""
        <div class="header">
            <h1>{content.title}</h1>
            <div class="subtitle">Generated on {content.generated_at.strftime('%B %d, %Y at %I:%M %p')}</div>
        </div>
        """)
        
        # Executive Summary
        html_parts.append(f"""
        <div class="executive-summary">
            <h2>Executive Summary</h2>
            <p>{content.executive_summary.replace(chr(10), '<br>')}</p>
        </div>
        """)
        
        # Process sections
        for section in content.sections:
            html_parts.append(await self._render_section(section))
        
        # Charts
        if content.charts:
            chart_generator = ChartGenerator()
            for chart in content.charts:
                chart_image = await chart_generator.generate_chart(chart)
                html_parts.append(f"""
                <div class="chart-container">
                    <h3>{chart.get('title', 'Chart')}</h3>
                    <img src="{chart_image}" alt="{chart.get('title', 'Chart')}" />
                </div>
                """)
        
        # Recommendations
        if content.recommendations:
            html_parts.append(await self._render_recommendations(content.recommendations))
        
        # Footer
        html_parts.append(f"""
        <div class="footer">
            <div>AITM Security Analysis Platform</div>
            <div class="timestamp">Report generated on {content.generated_at.strftime('%Y-%m-%d %H:%M:%S UTC')}</div>
        </div>
        """)
        
        return "".join(html_parts)
    
    async def _render_section(self, section: Dict[str, Any]) -> str:
        """Render a content section"""
        section_type = section.get("type", "generic")
        title = section.get("title", "Section")
        content_data = section.get("content", {})
        
        html = f'<div class="content-section"><h2 class="section-title">{title}</h2>'
        
        if section_type == "metrics":
            html += await self._render_metrics_section(content_data)
        elif section_type == "findings":
            html += await self._render_findings_section(content_data)
        elif section_type == "recommendations":
            html += await self._render_recommendations_section(content_data)
        else:
            html += await self._render_generic_section(content_data)
        
        html += '</div>'
        return html
    
    async def _render_metrics_section(self, content: Dict[str, Any]) -> str:
        """Render metrics as cards"""
        html = '<div class="metrics-grid">'
        
        for key, value in content.items():
            label = key.replace("_", " ").title()
            html += f"""
            <div class="metric-card">
                <div class="metric-value">{value}</div>
                <div class="metric-label">{label}</div>
            </div>
            """
        
        html += '</div>'
        return html
    
    async def _render_findings_section(self, content: List[Dict[str, Any]]) -> str:
        """Render findings list"""
        html = ""
        for finding in content:
            impact = finding.get("impact", "LOW").lower()
            html += f"""
            <div class="finding-item impact-{impact}">
                <h4>{finding.get("title", "Finding")}</h4>
                <p>{finding.get("description", "No description available.")}</p>
                <div><strong>Impact:</strong> {finding.get("impact", "Unknown")}</div>
                <div><strong>Trend:</strong> {finding.get("trend", "Unknown")}</div>
            </div>
            """
        return html
    
    async def _render_recommendations_section(self, content: List[Dict[str, Any]]) -> str:
        """Render recommendations list"""
        html = ""
        for rec in content:
            priority = rec.get("priority", "LOW").lower()
            html += f"""
            <div class="recommendation-item priority-{priority}">
                <h4>{rec.get("title", "Recommendation")}</h4>
                <p>{rec.get("description", "No description available.")}</p>
                <div><strong>Priority:</strong> {rec.get("priority", "Unknown")}</div>
                <div><strong>Timeline:</strong> {rec.get("timeline", "Unknown")}</div>
                <div><strong>Resources:</strong> {rec.get("resources", "Unknown")}</div>
            </div>
            """
        return html
    
    async def _render_generic_section(self, content: Any) -> str:
        """Render generic content"""
        if isinstance(content, dict):
            html = "<dl>"
            for key, value in content.items():
                label = key.replace("_", " ").title()
                if isinstance(value, (list, dict)):
                    value_str = json.dumps(value, indent=2) if value else "None"
                    html += f"<dt><strong>{label}:</strong></dt><dd><pre>{value_str}</pre></dd>"
                else:
                    html += f"<dt><strong>{label}:</strong></dt><dd>{value}</dd>"
            html += "</dl>"
            return html
        elif isinstance(content, list):
            html = "<ul>"
            for item in content:
                html += f"<li>{item}</li>"
            html += "</ul>"
            return html
        else:
            return f"<p>{content}</p>"
    
    async def _render_recommendations(self, recommendations: List[Dict[str, Any]]) -> str:
        """Render recommendations section"""
        if not recommendations:
            return ""
        
        html = '''
        <div class="recommendations">
            <h3>Strategic Recommendations</h3>
        '''
        
        for rec in recommendations:
            priority = rec.get("priority", "LOW").lower()
            html += f"""
            <div class="recommendation-item priority-{priority}">
                <h4>{rec.get("title", "Recommendation")}</h4>
                <p>{rec.get("description", "No description available.")}</p>
                <div><strong>Business Impact:</strong> {rec.get("business_impact", "Unknown")}</div>
                <div><strong>Timeline:</strong> {rec.get("timeline", "Unknown")}</div>
                <div><strong>Resources:</strong> {rec.get("resources", "Unknown")}</div>
            </div>
            """
        
        html += "</div>"
        return html


class PDFGenerator:
    """PDF generation service using WeasyPrint"""
    
    def __init__(self):
        self.html_engine = HTMLTemplateEngine()
        self.font_config = FontConfiguration() if WEASYPRINT_AVAILABLE else None
    
    async def generate_pdf(self, content: ReportContent) -> bytes:
        """Generate PDF from report content"""
        if not WEASYPRINT_AVAILABLE:
            raise ValueError("PDF generation not available. WeasyPrint not installed.")
        
        # Generate HTML
        html_content = await self.html_engine.render_report(content)
        
        # Create PDF
        html_doc = HTML(string=html_content)
        css_styles = CSS(string=self._get_pdf_css())
        
        pdf_bytes = html_doc.write_pdf(
            stylesheets=[css_styles],
            font_config=self.font_config
        )
        
        return pdf_bytes
    
    def _get_pdf_css(self) -> str:
        """Get PDF-specific CSS styles"""
        return """
        @page {
            size: A4;
            margin: 1in;
            @bottom-right {
                content: "Page " counter(page) " of " counter(pages);
                font-size: 10pt;
                color: #666;
            }
        }
        body {
            font-size: 12pt;
            line-height: 1.4;
        }
        .header {
            page-break-after: avoid;
        }
        .content-section {
            page-break-inside: avoid;
            margin-bottom: 20pt;
        }
        .chart-container {
            page-break-inside: avoid;
        }
        h1, h2, h3 {
            page-break-after: avoid;
        }
        """


class DOCXGenerator:
    """DOCX generation service using python-docx"""
    
    def __init__(self):
        self.available = DOCX_AVAILABLE
    
    async def generate_docx(self, content: ReportContent) -> bytes:
        """Generate DOCX from report content"""
        if not self.available:
            raise ValueError("DOCX generation not available. python-docx not installed.")
        
        doc = Document()
        
        # Set document styles
        self._setup_document_styles(doc)
        
        # Add title
        title = doc.add_heading(content.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_para = doc.add_paragraph(f"Generated on {content.generated_at.strftime('%B %d, %Y at %I:%M %p')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(content.executive_summary)
        
        # Process sections
        for section in content.sections:
            await self._add_section_to_doc(doc, section)
        
        # Add recommendations
        if content.recommendations:
            doc.add_heading('Strategic Recommendations', 1)
            for i, rec in enumerate(content.recommendations, 1):
                doc.add_heading(f"{i}. {rec.get('title', 'Recommendation')}", 2)
                doc.add_paragraph(rec.get('description', 'No description available.'))
                
                details_table = doc.add_table(rows=3, cols=2)
                details_table.style = 'Table Grid'
                
                details_table.cell(0, 0).text = "Priority"
                details_table.cell(0, 1).text = rec.get('priority', 'Unknown')
                details_table.cell(1, 0).text = "Timeline"
                details_table.cell(1, 1).text = rec.get('timeline', 'Unknown')
                details_table.cell(2, 0).text = "Resources"
                details_table.cell(2, 1).text = rec.get('resources', 'Unknown')
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _setup_document_styles(self, doc):
        """Setup document styles"""
        styles = doc.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.color.rgb = RGBColor(0, 0, 0)
        title_font.bold = True
    
    async def _add_section_to_doc(self, doc, section: Dict[str, Any]):
        """Add a section to the document"""
        title = section.get("title", "Section")
        content_data = section.get("content", {})
        section_type = section.get("type", "generic")
        
        doc.add_heading(title, 1)
        
        if section_type == "metrics":
            await self._add_metrics_to_doc(doc, content_data)
        elif section_type == "findings":
            await self._add_findings_to_doc(doc, content_data)
        elif section_type == "recommendations":
            await self._add_recommendations_to_doc(doc, content_data)
        else:
            await self._add_generic_content_to_doc(doc, content_data)
    
    async def _add_metrics_to_doc(self, doc, content: Dict[str, Any]):
        """Add metrics table to document"""
        if not content:
            return
        
        table = doc.add_table(rows=len(content), cols=2)
        table.style = 'Table Grid'
        
        for i, (key, value) in enumerate(content.items()):
            table.cell(i, 0).text = key.replace("_", " ").title()
            table.cell(i, 1).text = str(value)
    
    async def _add_findings_to_doc(self, doc, content: List[Dict[str, Any]]):
        """Add findings to document"""
        for i, finding in enumerate(content, 1):
            doc.add_heading(f"{i}. {finding.get('title', 'Finding')}", 2)
            doc.add_paragraph(finding.get('description', 'No description available.'))
            
            details_para = doc.add_paragraph()
            details_para.add_run(f"Impact: ").bold = True
            details_para.add_run(f"{finding.get('impact', 'Unknown')}\n")
            details_para.add_run(f"Trend: ").bold = True
            details_para.add_run(f"{finding.get('trend', 'Unknown')}")
    
    async def _add_recommendations_to_doc(self, doc, content: List[Dict[str, Any]]):
        """Add recommendations to document"""
        for i, rec in enumerate(content, 1):
            doc.add_heading(f"{i}. {rec.get('title', 'Recommendation')}", 2)
            doc.add_paragraph(rec.get('description', 'No description available.'))
    
    async def _add_generic_content_to_doc(self, doc, content: Any):
        """Add generic content to document"""
        if isinstance(content, dict):
            for key, value in content.items():
                para = doc.add_paragraph()
                para.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                para.add_run(str(value))
        elif isinstance(content, list):
            for item in content:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph(str(content))


class ReportFormatter:
    """Main report formatting service"""
    
    def __init__(self):
        self.html_engine = HTMLTemplateEngine()
        self.pdf_generator = PDFGenerator()
        self.docx_generator = DOCXGenerator()
        self.chart_generator = ChartGenerator()
    
    async def format_report(self, content: ReportContent, format_type: ReportFormat) -> Union[str, bytes]:
        """Format report content to specified format"""
        logger.info(f"Formatting report: {content.title} as {format_type}")
        
        if format_type == ReportFormat.HTML:
            return await self.html_engine.render_report(content)
        elif format_type == ReportFormat.PDF:
            return await self.pdf_generator.generate_pdf(content)
        elif format_type == ReportFormat.DOCX:
            return await self.docx_generator.generate_docx(content)
        elif format_type == ReportFormat.JSON:
            return await self._format_as_json(content)
        elif format_type == ReportFormat.MARKDOWN:
            return await self._format_as_markdown(content)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    async def _format_as_json(self, content: ReportContent) -> str:
        """Format as JSON"""
        data = {
            "title": content.title,
            "executive_summary": content.executive_summary,
            "sections": content.sections,
            "charts": content.charts,
            "recommendations": content.recommendations,
            "metadata": content.metadata,
            "generated_at": content.generated_at.isoformat()
        }
        return json.dumps(data, indent=2, default=str)
    
    async def _format_as_markdown(self, content: ReportContent) -> str:
        """Format as Markdown"""
        md_parts = []
        
        # Title and metadata
        md_parts.append(f"# {content.title}\n")
        md_parts.append(f"*Generated on {content.generated_at.strftime('%B %d, %Y at %I:%M %p')}*\n")
        
        # Executive Summary
        md_parts.append("## Executive Summary\n")
        md_parts.append(f"{content.executive_summary}\n")
        
        # Sections
        for section in content.sections:
            title = section.get("title", "Section")
            md_parts.append(f"## {title}\n")
            
            content_data = section.get("content", {})
            if isinstance(content_data, dict):
                for key, value in content_data.items():
                    md_parts.append(f"**{key.replace('_', ' ').title()}:** {value}\n")
            elif isinstance(content_data, list):
                for item in content_data:
                    if isinstance(item, dict):
                        item_title = item.get("title", "Item")
                        md_parts.append(f"### {item_title}\n")
                        if "description" in item:
                            md_parts.append(f"{item['description']}\n")
                    else:
                        md_parts.append(f"- {item}\n")
            else:
                md_parts.append(f"{content_data}\n")
            
            md_parts.append("\n")
        
        # Recommendations
        if content.recommendations:
            md_parts.append("## Strategic Recommendations\n")
            for i, rec in enumerate(content.recommendations, 1):
                md_parts.append(f"### {i}. {rec.get('title', 'Recommendation')}\n")
                md_parts.append(f"{rec.get('description', 'No description available.')}\n")
                md_parts.append(f"**Priority:** {rec.get('priority', 'Unknown')}\n")
                md_parts.append(f"**Timeline:** {rec.get('timeline', 'Unknown')}\n")
                md_parts.append(f"**Resources:** {rec.get('resources', 'Unknown')}\n\n")
        
        return "\n".join(md_parts)
    
    async def get_supported_formats(self) -> List[ReportFormat]:
        """Get list of supported formats"""
        formats = [ReportFormat.HTML, ReportFormat.JSON, ReportFormat.MARKDOWN]
        
        if WEASYPRINT_AVAILABLE:
            formats.append(ReportFormat.PDF)
        if DOCX_AVAILABLE:
            formats.append(ReportFormat.DOCX)
            
        return formats


# Global formatter instance
report_formatter = ReportFormatter()


async def format_report_async(content: ReportContent, format_type: ReportFormat) -> Union[str, bytes]:
    """Async wrapper for report formatting"""
    return await report_formatter.format_report(content, format_type)
